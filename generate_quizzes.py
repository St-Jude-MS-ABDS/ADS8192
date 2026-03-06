"""
Generate .docx quiz files for ADS 8192 Unit 1.
Each quiz: 4-5 questions (MC, select-all, short answer) focused on
theory/design/rationale, with answer key after a page break.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

ASSESSMENTS_DIR = os.path.join(os.path.dirname(__file__), "Unit1", "Assessments")


def add_page_break(doc):
    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_question(doc, num, qtype, text):
    """Add a formatted question."""
    p = doc.add_paragraph()
    run = p.add_run(f"Question {num} ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(f"({qtype})")
    run.italic = True
    run.font.size = Pt(11)
    p.space_after = Pt(4)

    p2 = doc.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(6)
    return p2


def add_choice(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{letter})  ")
    run.bold = True
    p.add_run(text)
    return p


def add_blank_lines(doc, n=2):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def add_answer(doc, num, answer_text):
    p = doc.add_paragraph()
    run = p.add_run(f"Question {num}: ")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(answer_text)
    p.paragraph_format.space_after = Pt(8)
    return p


def make_quiz(filename, title, questions, answers):
    """
    questions: list of dicts with keys: num, type, text, choices (optional)
    answers: list of dicts with keys: num, text
    """
    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Course info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ADS 8192 — Developing Scientific Applications")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # Instructions
    p = doc.add_paragraph()
    run = p.add_run("Instructions: ")
    run.bold = True
    p.add_run(
        "Answer all questions. For multiple choice, select the single best answer. "
        "For select-all-that-apply, select every correct option. "
        "For short answer, respond in 1–3 sentences."
    )
    p.paragraph_format.space_after = Pt(12)

    # Separator
    doc.add_paragraph("─" * 60)

    # Questions
    for q in questions:
        add_question(doc, q["num"], q["type"], q["text"])
        if "choices" in q:
            for letter, choice_text in q["choices"]:
                add_choice(doc, letter, choice_text)
        if q["type"] == "Short Answer":
            add_blank_lines(doc, 3)
        doc.add_paragraph()  # spacer between questions

    # Page break before answer key
    add_page_break(doc)

    add_heading(doc, "Answer Key", level=1)
    doc.add_paragraph("─" * 60)

    for a in answers:
        add_answer(doc, a["num"], a["text"])

    # Save
    filepath = os.path.join(ASSESSMENTS_DIR, filename)
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ─── Quiz 03 ───────────────────────────────────────────────

def quiz03():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "What is the PRIMARY design advantage of storing gene expression data, sample metadata, and gene annotations together in a single container (like SummarizedExperiment) rather than as separate objects?",
            "choices": [
                ("A", "It uses less memory than separate objects"),
                ("B", "It automatically normalizes the data"),
                ("C", "Subsetting or reordering keeps all components synchronized, preventing mismatches"),
                ("D", "It makes the data compatible with Python"),
            ],
        },
        {
            "num": 2,
            "type": "Select All That Apply",
            "text": "Which of the following are good design principles for S4 classes in Bioconductor? (Select all that apply)",
            "choices": [
                ("A", "Provide constructor functions so users don't have to call new() directly"),
                ("B", "Include validity checks to catch invalid data early"),
                ("C", "Store all derived values in slots, even if they can be recomputed"),
                ("D", "Use accessor methods instead of direct slot access with @"),
            ],
        },
        {
            "num": 3,
            "type": "Multiple Choice",
            "text": 'In the "analysis core" design pattern, functions like run_pca() accept a SummarizedExperiment and return structured output. Why is this approach better than writing one long script that performs the entire analysis?',
            "choices": [
                ("A", "Long scripts run slower because R interprets them line by line"),
                ("B", "Small, composable functions can be independently tested, reused across interfaces (Shiny, CLI), and debugged in isolation"),
                ("C", "R does not support scripts longer than 100 lines"),
                ("D", "Functions automatically parallelize across CPU cores"),
            ],
        },
        {
            "num": 4,
            "type": "Short Answer",
            "text": "A function returns only a matrix of filtered gene expression values. A colleague argues it should return a subsetted SummarizedExperiment instead. In 1–2 sentences, explain why the colleague's approach is better.",
        },
        {
            "num": 5,
            "type": "Multiple Choice",
            "text": "When designing a package, why is it better for functions to accept and return structured containers (like SummarizedExperiment) rather than requiring users to pass multiple separate arguments (counts matrix, sample table, gene table)?",
            "choices": [
                ("A", "Structured containers use less memory than separate objects"),
                ("B", "It reduces the number of arguments, prevents mismatches between components, and makes functions composable—the output of one can be passed directly to the next"),
                ("C", "R does not allow functions to have more than three arguments"),
                ("D", "Separate arguments are only supported in Python, not R"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "C. When you subset a SummarizedExperiment (e.g., se[1:100, 1:5]), the assays, rowData, and colData are all subsetted together automatically. This prevents the common and dangerous bug of having mismatched data and metadata."},
        {"num": 2, "text": "A, B, D. Constructors (A) ensure consistent creation, validity checks (B) catch errors early, and accessors (D) provide a stable interface. Storing derived values (C) is discouraged because it introduces redundancy and risks inconsistency."},
        {"num": 3, "text": "B. Small functions with clear inputs and outputs are testable in isolation, reusable across different interfaces (R console, Shiny, CLI), and easier to debug. This is the foundation of the 'three interfaces, one core' architecture."},
        {"num": 4, "text": "Returning a subsetted SummarizedExperiment preserves the connection between filtered expression values and their gene annotations and sample metadata. A bare matrix loses that context, making downstream analysis error-prone and requiring users to manually track which genes and samples remain."},
        {"num": 5, "text": "B. Structured containers bundle related data together, so functions need fewer arguments and there's no risk of passing a counts matrix that doesn't match the sample metadata. Functions that return the same container type are naturally composable—you can chain top_variable_features() into run_pca() without manually juggling multiple objects."},
    ]

    make_quiz("Quiz03_Data_Structures_Bioconductor.docx", "Quiz 3: Data Structures & Bioconductor", questions, answers)


# ─── Quiz 05 ───────────────────────────────────────────────

def quiz05():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "What is the fundamental purpose of the DESCRIPTION file in an R package?",
            "choices": [
                ("A", "To store the source code for all package functions"),
                ("B", "To declare what the package is, who wrote it, what it depends on, and how it's licensed"),
                ("C", "To configure the testing framework"),
                ("D", "To list every file that should be included in the package"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "Why should you NEVER use library() inside an R package's source code?",
            "choices": [
                ("A", "library() is slower than require()"),
                ("B", "library() modifies the global search path, which can cause conflicts and makes dependencies implicit rather than declared"),
                ("C", "library() only works in RStudio, not in R CMD check"),
                ("D", "library() doesn't work with Bioconductor packages"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": "Which of the following describe the role of the NAMESPACE file? (Select all that apply)",
            "choices": [
                ("A", "It controls which functions are visible to users when they load the package"),
                ("B", "It declares which functions from other packages your code uses"),
                ("C", "It stores the package version number"),
                ("D", "It is generated automatically by roxygen2 from your documentation comments"),
            ],
        },
        {
            "num": 4,
            "type": "Multiple Choice",
            "text": "Your package uses testthat for testing but never calls testthat functions in its core code. Where should testthat be listed in DESCRIPTION?",
            "choices": [
                ("A", "Imports, because every package that uses testthat must declare it there"),
                ("B", "Suggests, because it is only needed by developers running tests—not by users of the package"),
                ("C", "Depends, so that testthat is always loaded when your package is loaded"),
                ("D", "It should not appear in DESCRIPTION at all"),
            ],
        },
        {
            "num": 5,
            "type": "Multiple Choice",
            "text": "What is the key benefit of the roxygen2 documentation workflow (writing documentation as comments above functions)?",
            "choices": [
                ("A", "It makes your code run faster"),
                ("B", "It keeps documentation next to the code it describes, so they stay in sync as the code changes"),
                ("C", "It removes the need for a DESCRIPTION file"),
                ("D", "It automatically translates documentation into multiple languages"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. The DESCRIPTION file is the package's identity card—it contains metadata (name, title, version, authors), declares dependencies (Imports, Suggests), specifies the license, and is read by R's package infrastructure for installation and checking."},
        {"num": 2, "text": "B. library() attaches a package to the global search path, which can mask functions from other packages and creates implicit dependencies. Inside a package, you should use explicit namespacing (pkg::fun()) or roxygen2 @import directives, with dependencies declared in DESCRIPTION."},
        {"num": 3, "text": "A, B, D. The NAMESPACE controls exports (what users see, A) and imports (what your code uses from other packages, B). It is generated by roxygen2 from @export and @import tags (D). The version number lives in DESCRIPTION, not NAMESPACE (C)."},
        {"num": 4, "text": "B. Suggests packages are not installed automatically with your package, so end users who just want to use the core functions don't have to install testthat. It belongs in Suggests because only developers running tests need it."},
        {"num": 5, "text": "B. When documentation lives right above the function it describes, developers naturally update both together. The alternative—maintaining separate .Rd files—leads to documentation that drifts out of sync with the code."},
    ]

    make_quiz("Quiz05_Package_Development_devtools.docx", "Quiz 5: R Package Development (devtools)", questions, answers)


# ─── Quiz 06 ───────────────────────────────────────────────

def quiz06():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "What is the PRIMARY purpose of automated testing in scientific software?",
            "choices": [
                ("A", "To prove the software is mathematically correct"),
                ("B", "To catch regressions—ensuring that changes don't break existing functionality"),
                ("C", "To replace manual code review"),
                ("D", "To make the software run faster"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "Why is continuous integration (CI) valuable even if all your tests pass on your own machine?",
            "choices": [
                ("A", "CI servers have faster processors"),
                ("B", "CI runs tests in a clean environment, catching undeclared dependencies and platform-specific issues"),
                ("C", "CI automatically fixes bugs before they reach users"),
                ("D", "CI is required by all R package repositories"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": "Which of the following are good strategies when writing tests for analysis functions? (Select all that apply)",
            "choices": [
                ("A", "Test edge cases like empty inputs, single-element inputs, and missing values"),
                ("B", "Verify that output dimensions and types match expectations"),
                ("C", "Test only the 'happy path'—assume users always provide correct inputs"),
                ("D", "Write the test before or alongside the function to clarify expected behavior"),
            ],
        },
        {
            "num": 4,
            "type": "Short Answer",
            "text": "A colleague's package has no documentation website. They ask why they should bother setting up pkgdown. In 1–2 sentences, explain the main benefit.",
        },
        {
            "num": 5,
            "type": "Multiple Choice",
            "text": "In the context of test-driven development, what is the recommended workflow?",
            "choices": [
                ("A", "Write all the code first, then write tests at the end"),
                ("B", "Write a test that fails, write the minimum code to make it pass, then refactor"),
                ("C", "Write tests only for functions that have known bugs"),
                ("D", "Let users report bugs, then write tests for those specific cases"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. Tests act as a safety net: when you modify or extend code, they verify that everything that worked before still works. This is especially important in scientific software where subtle numerical changes can invalidate results."},
        {"num": 2, "text": "B. Your local machine may have extra packages installed, specific OS configurations, or cached files that mask problems. CI starts from a clean slate with only declared dependencies, revealing issues other users will encounter."},
        {"num": 3, "text": "A, B, D. Edge cases (A) catch common failure modes, checking output structure (B) verifies contracts, and writing tests first (D) clarifies design intent. Testing only the happy path (C) leaves the most common sources of bugs untested."},
        {"num": 4, "text": "pkgdown generates a browsable, searchable documentation website from existing help files and vignettes—with no extra writing required. This makes the package accessible to users who prefer reading documentation in a web browser rather than the R help viewer."},
        {"num": 5, "text": "B. The red-green-refactor cycle (fail → pass → clean up) ensures every piece of code exists to satisfy a clear requirement, avoids writing untested code, and produces a thorough test suite as a natural byproduct of development."},
    ]

    make_quiz("Quiz06_pkgdown_testthat.docx", "Quiz 6: Testing & Documentation (pkgdown / testthat)", questions, answers)


# ─── Quiz 07 ───────────────────────────────────────────────

def quiz07():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "What is the fundamental difference between a reactive expression and an observer in Shiny?",
            "choices": [
                ("A", "Reactive expressions are faster than observers"),
                ("B", "A reactive expression produces a cached value that other code can consume; an observer performs side effects and returns nothing"),
                ("C", "Observers run on the client; reactive expressions run on the server"),
                ("D", "There is no practical difference; they are interchangeable"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "An expensive computation (e.g., PCA) is placed directly inside renderPlot(). The same data is also needed by a summary table. What is the design problem, and what is the solution?",
            "choices": [
                ("A", "No problem—Shiny automatically shares results between outputs"),
                ("B", "The computation runs separately for each output; extract it into a reactive() so it computes once and is shared"),
                ("C", "The computation should be placed in the UI function instead"),
                ("D", "You should save results to a file and read them from each output"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": "What problems do Shiny modules solve? (Select all that apply)",
            "choices": [
                ("A", "They prevent input/output ID collisions when the same UI component is reused"),
                ("B", "They bundle related UI and server logic into a reusable, self-contained unit"),
                ("C", "They automatically make the app load faster"),
                ("D", "They make it easier to test pieces of the app in isolation"),
            ],
        },
        {
            "num": 4,
            "type": "Multiple Choice",
            "text": "Shiny's reactive programming model is described as 'declarative' rather than 'imperative.' What does this mean in practice?",
            "choices": [
                ("A", "You must write explicit update calls every time an input changes"),
                ("B", "You describe relationships between values and Shiny automatically determines what to recompute and when"),
                ("C", "All outputs are computed once at app startup and never change"),
                ("D", "Declarative and imperative are synonyms in the Shiny context"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. A reactive() is like a formula cell in a spreadsheet—it computes and caches a value that downstream code references. An observe() is for actions (logging, writing files, updating UI) that don't produce a reusable value."},
        {"num": 2, "text": "B. Placing the computation in renderPlot() means it re-runs for every output that needs it. Wrapping it in reactive() computes it once; both the plot and the table then call pca_result() to get the cached value."},
        {"num": 3, "text": "A, B, D. Modules use namespaced IDs to prevent collisions (A), encapsulate UI + server logic (B), and enable isolated testing (D). They don't inherently speed up the app (C)—performance depends on reactive design."},
        {"num": 4, "text": "B. In declarative programming, you describe relationships (e.g., 'this plot depends on these inputs') rather than writing step-by-step instructions for when to update. Shiny's reactive graph automatically determines what to recompute and when, so you focus on what depends on what—not the order of execution."},
    ]

    make_quiz("Quiz07_Shiny_Reactivity_Modules.docx", "Quiz 7: Shiny Reactivity & Modules", questions, answers)


# ─── Quiz 08 ───────────────────────────────────────────────

def quiz08():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "Why is it valuable to package a Shiny app inside an R package rather than distributing it as standalone app.R files?",
            "choices": [
                ("A", "Shiny apps cannot run outside of R packages"),
                ("B", "Packaging provides dependency management, versioning, documentation, and a standard installation workflow"),
                ("C", "R packages make Shiny apps run faster"),
                ("D", "CRAN requires all Shiny apps to be in packages"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "The inst/ directory has a special behavior during R package installation. What happens to its contents?",
            "choices": [
                ("A", "They are compiled into binary code"),
                ("B", "They are deleted to reduce package size"),
                ("C", "They are copied to the installed package root, with the inst/ prefix removed"),
                ("D", "They are moved to the R/ directory automatically"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": 'In the "three interfaces, one core" architecture, which of the following are true? (Select all that apply)',
            "choices": [
                ("A", "The Shiny app, CLI, and direct R usage all call the same underlying analysis functions"),
                ("B", "Each interface has its own copy of the analysis logic for maximum performance"),
                ("C", "Fixing a bug in a core function automatically fixes it across all three interfaces"),
                ("D", "The architecture makes it easier to test the analysis logic independently of any interface"),
            ],
        },
        {
            "num": 4,
            "type": "True/False",
            "text": "True or False: system.file() should be used to locate package files at runtime because it resolves the correct installed path regardless of platform, R version, or library location.",
            "choices": [
                ("A", "True"),
                ("B", "False"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. A package brings formal dependency declaration (DESCRIPTION), semantic versioning, standard installation (install.packages / remotes::install_github), automated testing, and documentation. Users get a reliable one-command installation experience instead of manually sourcing scripts."},
        {"num": 2, "text": "C. During installation, the inst/ prefix is stripped—so inst/shiny/app.R becomes accessible at the package root as shiny/app.R. You locate it at runtime with system.file('shiny', 'app.R', package = 'yourpkg')."},
        {"num": 3, "text": "A, C, D. The whole point is a shared core: all interfaces call the same functions (A), so a bug fix propagates everywhere (C), and the core can be tested without spinning up Shiny or a CLI (D). Having separate copies (B) defeats the purpose."},
        {"num": 4, "text": "A (True). system.file() resolves the installed location of package files at runtime, which varies by platform, R version, and library path. Hardcoded paths break as soon as the package is installed somewhere other than where you developed it."},
    ]

    make_quiz("Quiz08_Shiny_Packaging_Deployment.docx", "Quiz 8: Shiny App Packaging & Deployment", questions, answers)


# ─── Quiz 09 ───────────────────────────────────────────────

def quiz09():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "Why are command-line interfaces important for scientific software, even when a Shiny app already exists?",
            "choices": [
                ("A", "CLIs are always faster than graphical interfaces"),
                ("B", "CLIs enable automation, reproducibility, and integration into computational pipelines and HPC environments"),
                ("C", "Shiny apps cannot process real data"),
                ("D", "CLIs are required by all scientific journals"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "What is the purpose of using distinct exit codes (0 for success, non-zero for errors) in a CLI tool?",
            "choices": [
                ("A", "They control how much memory the program uses"),
                ("B", "They allow calling scripts and pipelines to programmatically detect whether the tool succeeded or failed"),
                ("C", "They determine the order in which output lines are printed"),
                ("D", "They are only meaningful on Windows, not Unix systems"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": "Which of the following are principles of well-designed CLI tools? (Select all that apply)",
            "choices": [
                ("A", "Write error and diagnostic messages to stderr, results to stdout"),
                ("B", "Provide a --help flag that documents all arguments"),
                ("C", "Support machine-readable output formats (e.g., CSV) for downstream scripting"),
                ("D", "Require interactive user input for every parameter"),
            ],
        },
        {
            "num": 4,
            "type": "Short Answer",
            "text": "The DRY principle ('Don't Repeat Yourself') is central to CLI design in this course. In 1–2 sentences, explain what it means and why it matters when adding a CLI to an existing package.",
        },
    ]

    answers = [
        {"num": 1, "text": "B. CLIs can be called from shell scripts, workflow managers (Snakemake, Nextflow), and HPC schedulers (SLURM). This enables automated, reproducible analyses at scale—something an interactive GUI cannot provide."},
        {"num": 2, "text": "B. Scripts and pipelines check exit codes to decide what to do next (e.g., stop on failure, retry, log the error). Without meaningful exit codes, automated workflows can't distinguish success from failure."},
        {"num": 3, "text": "A, B, C. Separating stderr/stdout (A) enables piping, --help (B) makes the tool self-documenting, and machine-readable output (C) supports scripting. Requiring interactive input (D) defeats the purpose of a CLI—it should work non-interactively with arguments."},
        {"num": 4, "text": "DRY means the CLI should call the same core functions as the Shiny app and direct R usage—not duplicate their logic. This ensures consistency across interfaces: fix a bug once and it's fixed everywhere, rather than maintaining parallel implementations that can diverge."},
    ]

    make_quiz("Quiz09_CLI_Design.docx", "Quiz 9: Command-Line Interface Design", questions, answers)


# ─── Quiz 10 ───────────────────────────────────────────────

def quiz10():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "What is the purpose of 'clean-room' testing when packaging a CLI tool?",
            "choices": [
                ("A", "To test the tool in a physically clean laboratory environment"),
                ("B", "To verify the tool works in a fresh R installation with only declared dependencies, catching hidden assumptions"),
                ("C", "To remove all test files before releasing the package"),
                ("D", "To ensure the tool only runs on the developer's machine"),
            ],
        },
        {
            "num": 2,
            "type": "Multiple Choice",
            "text": "What does 'output parity' mean in the context of a CLI wrapper around package functions?",
            "choices": [
                ("A", "The CLI must produce output that looks identical on all operating systems"),
                ("B", "The CLI and the underlying R functions must produce the same numerical results and handle errors the same way"),
                ("C", "The CLI output must be an even number of bytes"),
                ("D", "The CLI must print its output in pairs of lines"),
            ],
        },
        {
            "num": 3,
            "type": "Select All That Apply",
            "text": "Which of the following should be part of a package release checklist? (Select all that apply)",
            "choices": [
                ("A", "R CMD check passes with no errors, warnings, or notes"),
                ("B", "All tests pass in a clean environment, not just on your machine"),
                ("C", "Documentation is up to date and roxygen2 has been run"),
                ("D", "The package has been renamed to a new name for each release"),
            ],
        },
        {
            "num": 4,
            "type": "Multiple Choice",
            "text": "After a user installs your package from GitHub, the exec/ directory containing your Rapp CLI is buried inside the R library tree. What is the recommended way to let users invoke the CLI directly from a terminal (e.g., sePCA pca --help)?",
            "choices": [
                ("A", "Tell users to find the library path manually and add it to their PATH"),
                ("B", "Export an install helper function (e.g., install_sePCA_cli()) that calls Rapp::install_pkg_cli_apps() to place lightweight launcher scripts on the user's PATH"),
                ("C", "Copy the exec/ script to the user's Desktop during package installation"),
                ("D", "Require users to always run the CLI via Rscript -e instead of a direct command"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. Clean-room testing installs the package in a fresh library with no extra packages. This catches undeclared dependencies, missing imports, and assumptions that only hold on your development machine."},
        {"num": 2, "text": "B. Output parity means the CLI wrapper is just a thin layer around the package functions—same inputs produce the same results, same invalid inputs produce the same errors. This is a direct consequence of the DRY / 'three interfaces, one core' design."},
        {"num": 3, "text": "A, B, C. R CMD check (A) is the standard validation tool, clean-environment testing (B) catches hidden dependencies, and updated documentation (C) ensures users can use the package correctly. Renaming the package (D) would break everything."},
        {"num": 4, "text": "B. Rapp::install_pkg_cli_apps() creates lightweight launcher scripts (.bat on Windows, shell scripts on Unix) that live on the user's PATH and forward to the installed Rapp app. Exporting a convenience wrapper like install_sePCA_cli() gives users a one-command setup experience."},
    ]

    make_quiz("Quiz10_CLI_Packaging_Installation.docx", "Quiz 10: CLI Packaging & Installation", questions, answers)


# ─── Quiz 11 ───────────────────────────────────────────────

def quiz11():
    questions = [
        {
            "num": 1,
            "type": "Multiple Choice",
            "text": "When you encounter an error in your R package, what is the recommended FIRST step in systematic debugging?",
            "choices": [
                ("A", "Rewrite the function from scratch"),
                ("B", "Read the error message carefully and identify exactly which function and input caused it"),
                ("C", "Install a different version of R"),
                ("D", "Delete the test that found the error"),
            ],
        },
        {
            "num": 2,
            "type": "Select All That Apply",
            "text": "What should a good reproducible bug report include? (Select all that apply)",
            "choices": [
                ("A", "A minimal example that demonstrates the problem"),
                ("B", "The exact error message or description of unexpected behavior"),
                ("C", "Your entire codebase so the reader can find the relevant parts"),
                ("D", "R version and relevant package versions (sessionInfo output)"),
            ],
        },
        {
            "num": 3,
            "type": "Multiple Choice",
            "text": "Your package passes R CMD check locally but a classmate cannot install it. What is the most likely explanation?",
            "choices": [
                ("A", "R packages only work on the machine where they were built"),
                ("B", "A dependency you use locally is not declared in DESCRIPTION, so it's missing on their machine"),
                ("C", "Your classmate needs to use the exact same operating system"),
                ("D", "R CMD check does not actually verify that the package works"),
            ],
        },
        {
            "num": 4,
            "type": "Multiple Choice",
            "text": "During peer review of a classmate's package, which of the following is the MOST useful type of feedback?",
            "choices": [
                ("A", "Telling them their code style is wrong"),
                ("B", "Identifying a specific scenario where the package produces incorrect or unexpected results"),
                ("C", "Suggesting they rewrite the entire package in Python"),
                ("D", "Pointing out that their variable names are different from yours"),
            ],
        },
    ]

    answers = [
        {"num": 1, "text": "B. The error message tells you what went wrong and often where. Understanding the problem before attempting a fix prevents wasted effort solving the wrong issue and helps you write a targeted test to prevent the bug from recurring."},
        {"num": 2, "text": "A, B, D. A minimal example (A) lets others reproduce the issue, the exact error (B) provides diagnostic information, and session info (D) catches version/platform issues. Sharing the entire codebase (C) makes the problem harder to isolate."},
        {"num": 3, "text": "B. If it works locally but not elsewhere, you almost certainly have packages installed that your code uses but that aren't declared in DESCRIPTION. Clean-room testing (a fresh library with only declared dependencies) catches this class of error."},
        {"num": 4, "text": "B. Actionable, specific feedback about incorrect behavior is the most valuable kind of review. It identifies a concrete problem the author can fix and test. Stylistic preferences and language suggestions are subjective and less impactful."},
    ]

    make_quiz("Quiz11_Review_Debugging.docx", "Quiz 11: Review & Debugging Workflows", questions, answers)


# ─── Main ──────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs(ASSESSMENTS_DIR, exist_ok=True)
    print("Generating quizzes...")
    quiz03()
    quiz05()
    quiz06()
    quiz07()
    quiz08()
    quiz09()
    quiz10()
    quiz11()
    print("Done! All quizzes generated.")
